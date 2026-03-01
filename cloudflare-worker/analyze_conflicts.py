#!/usr/bin/env python3
"""Analyze driver_manual.pdf and edited_support_services_sog.docx for conflicting SOG/policy content."""

import pdfplumber
from docx import Document

# Extract driver manual text
dm_pages = {}
with pdfplumber.open('/tmp/driver_manual.pdf') as pdf:
    for i, page in enumerate(pdf.pages):
        t = page.extract_text()
        if t:
            dm_pages[i+1] = t

# Extract SOG docx text  
doc = Document('/tmp/edited_support_services_sog.docx')
sog_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])

# SOG-related keywords that indicate policy/procedure content (not driver-specific technical content)
sog_keywords = ['sog', 'standard operating', 'policy', 'procedure', 'guidelines', 
                'support services', 'uniform', 'station supply', 'requisition', 'procurement',
                'inventory management', 'supply chain', 'ordering', 'distribution',
                'accountability', 'chain of command', 'responsibility', 'duty',
                'section', 'article', 'regulation', 'compliance', 'directive',
                'fire prevention', 'inspection schedule', 'maintenance schedule',
                'staffing', 'personnel', 'training requirement']

print('=== DRIVER MANUAL PAGES WITH SOG/POLICY CONTENT ===')
flagged_pages = []
for page_num, text in sorted(dm_pages.items()):
    text_lower = text.lower()
    matches = [kw for kw in sog_keywords if kw in text_lower]
    if matches:
        flagged_pages.append(page_num)
        preview = text[:500].replace('\n', ' | ')
        print(f'\n--- Page {page_num} (matched: {matches}) ---')
        print(preview)
        print('...')

print(f'\n\n=== TOTAL FLAGGED PAGES: {len(flagged_pages)} out of {len(dm_pages)} ===')
print(f'Pages: {flagged_pages}')

# Now look for specific SOG section headers/topics in the SOG docx
print('\n\n=== SOG DOCUMENT TOPIC HEADINGS ===')
for p in doc.paragraphs:
    text = p.text.strip()
    if text and (p.style.name.startswith('Heading') or text.isupper() or 
                 text.startswith('SOG') or text.startswith('Section') or
                 text.startswith('ARTICLE') or text.startswith('Policy')):
        print(f'  [{p.style.name}] {text[:150]}')

# Cross-reference: find driver manual pages that discuss topics covered in the SOG
print('\n\n=== CROSS-REFERENCE: Driver Manual pages likely containing outdated SOG info ===')
# Extract key topics from SOG
sog_topics = []
for p in doc.paragraphs:
    text = p.text.strip().lower()
    if len(text) > 15:
        sog_topics.append(text)

# Check each driver manual page for SOG topic overlap        
for page_num, text in sorted(dm_pages.items()):
    text_lower = text.lower()
    # Check for direct SOG references
    has_sog_ref = any(term in text_lower for term in ['sog ', 'sog#', 'standard operating guideline', 
                                                        'support services sog', 'operational guideline'])
    if has_sog_ref:
        print(f'\n  *** Page {page_num} DIRECTLY REFERENCES SOGs ***')
        # Find the relevant lines
        for line in text.split('\n'):
            line_lower = line.lower().strip()
            if any(term in line_lower for term in ['sog', 'standard operating', 'guideline', 'policy']):
                print(f'    > {line.strip()[:200]}')
